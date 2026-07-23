#!/usr/bin/env python3
"""Resolve and audit Phase 9 official numeric-target sources.

PDF text is extracted from every page, while expensive table recognition is limited to pages
that contain target-related vocabulary or have no extractable text. The module also provides the
shared source-resolution and extraction primitives used by the Reviewed catalog builder.
"""

from __future__ import annotations

import argparse
import hashlib
import io
import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from urllib.parse import urljoin, urlparse

import pdfplumber
import requests
from bs4 import BeautifulSoup
from docx import Document
from openpyxl import load_workbook

ROOT = Path(__file__).resolve().parents[1]
UPDATED_AT = "2026-07-22"
USER_AGENT = (
    "Mozilla/5.0 (compatible; JichiInsightEvidenceBot/1.0; "
    "+https://github.com/M-Osugi1230/jichi-insight)"
)
TARGET_KEYWORDS = (
    "数値目標",
    "目標値",
    "目標",
    "成果指標",
    "重要業績評価指標",
    "KPI",
    "ＫＰＩ",
    "指標",
    "達成目標",
    "政策目標",
    "施策目標",
    "現状値",
    "基準値",
)
LINK_KEYWORDS = (
    "数値",
    "目標",
    "指標",
    "kpi",
    "成果",
    "実施計画",
    "総合計画",
    "戦略",
    "進捗",
    "評価",
    "資料",
    "本編",
    "概要",
    "別冊",
    "付属",
    "附属",
)
DOCUMENT_SUFFIXES = (".pdf", ".xlsx", ".xlsm", ".docx")
NUMBER_RE = re.compile(r"(?<![A-Za-z])[▲△▽▼＋+－-]?[0-9０-９][0-9０-９,，.．％%]*")
YEAR_RE = re.compile(
    r"(?:令和|平成|昭和)?[0-9０-９]{1,4}(?:年度|年)|R[0-9]{1,2}|H[0-9]{1,2}|20[0-9]{2}"
)
SPACE_RE = re.compile(r"\s+")


@dataclass(frozen=True)
class ResolvedDocument:
    url: str
    title: str
    content_type: str
    score: int


def clean(value: str | None) -> str:
    if not value:
        return ""
    return SPACE_RE.sub(" ", value.replace("\u3000", " ").replace("\x00", " ")).strip()


def sha256_bytes(value: bytes) -> str:
    return hashlib.sha256(value).hexdigest()


def load_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def write_json(path: Path, value: object) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def source_registry_paths(root: Path) -> list[Path]:
    return sorted((root / "data/catalog").glob("phase9_*_source_registry.json"))


def official_phase9_records(root: Path) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    for path in source_registry_paths(root):
        registry = load_json(path)
        for record in registry["records"]:
            target_sources = [
                source for source in record["sources"] if source["role"] == "numeric_target_source"
            ]
            if len(target_sources) != 1:
                raise ValueError(
                    f"{path}: {record['prefecture_code']} must have one numeric target source"
                )
            records.append(
                {
                    "batch_id": registry["batch_id"],
                    "prefecture_code": record["prefecture_code"],
                    "name": record["name"],
                    "current_plan_title": record["current_plan_title"],
                    "current_plan_period": record["current_plan_period"],
                    "target_source_boundary": record["target_source_boundary"],
                    "target_source": target_sources[0],
                    "registry_path": str(path.relative_to(root)),
                }
            )
    records.sort(key=lambda item: item["prefecture_code"])
    if len(records) != 38:
        raise ValueError(f"Expected 38 Phase 9 prefectures, found {len(records)}")
    return records


def request(session: requests.Session, url: str) -> requests.Response:
    response = session.get(url, timeout=(12, 45), allow_redirects=True)
    response.raise_for_status()
    return response


def normalized_content_type(response: requests.Response) -> str:
    return response.headers.get("content-type", "").split(";", 1)[0].lower().strip()


def is_document_url(url: str) -> bool:
    return urlparse(url).path.lower().endswith(DOCUMENT_SUFFIXES)


def link_score(text: str, href: str, same_host: bool) -> int:
    combined = f"{text} {href}".lower()
    score = 2 if same_host else 0
    if is_document_url(href):
        score += 8
    for keyword in LINK_KEYWORDS:
        if keyword.lower() in combined:
            score += 2
    if any(token in combined for token in ("様式", "募集", "予算", "会議録", "議事録")):
        score -= 3
    if any(token in combined for token in ("過去", "旧計画", "バックナンバー")):
        score -= 2
    return score


def resolve_documents(
    session: requests.Session,
    source_url: str,
    source_title: str,
) -> tuple[list[ResolvedDocument], dict[str, Any]]:
    response = request(session, source_url)
    content_type = normalized_content_type(response)
    audit: dict[str, Any] = {
        "landing_final_url": response.url,
        "landing_status_code": response.status_code,
        "landing_content_type": content_type,
        "landing_sha256": sha256_bytes(response.content),
    }
    if is_document_url(response.url) or content_type in {
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }:
        return [ResolvedDocument(response.url, source_title, content_type, 100)], audit

    soup = BeautifulSoup(response.content, "html.parser")
    base_host = urlparse(response.url).netloc
    candidates: dict[str, ResolvedDocument] = {}
    for anchor in soup.find_all("a", href=True):
        href = urljoin(response.url, anchor.get("href", "").strip())
        if not href.startswith("http"):
            continue
        text = clean(anchor.get_text(" ", strip=True))
        score = link_score(text, href, urlparse(href).netloc == base_host)
        if score < 5:
            continue
        candidate = ResolvedDocument(href, text or source_title, "", score)
        previous = candidates.get(href)
        if previous is None or candidate.score > previous.score:
            candidates[href] = candidate

    ranked = sorted(candidates.values(), key=lambda item: (-item.score, item.url))
    direct_documents = [item for item in ranked if is_document_url(item.url)]
    selected = direct_documents[:6] or ranked[:4]
    if not selected:
        selected = [ResolvedDocument(response.url, source_title, content_type, 1)]
    audit["discovered_link_count"] = len(candidates)
    audit["selected_document_count"] = len(selected)
    return selected, audit


def target_context(text: str) -> bool:
    lowered = text.lower()
    return any(keyword.lower() in lowered for keyword in TARGET_KEYWORDS)


def pdf_rows(content: bytes) -> tuple[list[dict[str, Any]], int]:
    rows: list[dict[str, Any]] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        page_count = len(pdf.pages)
        for page_number, page in enumerate(pdf.pages, 1):
            text = page.extract_text(x_tolerance=2, y_tolerance=3) or ""
            lines = [clean(raw_line) for raw_line in text.splitlines()]
            lines = [line for line in lines if line]
            for line_number, line in enumerate(lines, 1):
                rows.append(
                    {
                        "location_kind": "pdf_text_line",
                        "page": page_number,
                        "row": line_number,
                        "text": line,
                    }
                )
            should_extract_tables = not lines or any(target_context(line) for line in lines)
            if not should_extract_tables:
                continue
            for table_number, table in enumerate(page.extract_tables(), 1):
                for row_number, raw_row in enumerate(table or [], 1):
                    cells = [clean(cell) for cell in (raw_row or [])]
                    cells = [cell for cell in cells if cell]
                    if cells:
                        rows.append(
                            {
                                "location_kind": "pdf_table_row",
                                "page": page_number,
                                "table": table_number,
                                "row": row_number,
                                "text": " | ".join(cells),
                                "cells": cells,
                            }
                        )
    return rows, page_count


def workbook_rows(content: bytes) -> tuple[list[dict[str, Any]], int]:
    rows: list[dict[str, Any]] = []
    workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    for sheet in workbook.worksheets:
        empty_streak = 0
        for row_number, raw_row in enumerate(sheet.iter_rows(values_only=True), 1):
            cells = [clean(str(cell)) if cell is not None else "" for cell in raw_row]
            cells = [cell for cell in cells if cell]
            if not cells:
                empty_streak += 1
                if empty_streak >= 100:
                    break
                continue
            empty_streak = 0
            rows.append(
                {
                    "location_kind": "workbook_row",
                    "sheet": sheet.title,
                    "row": row_number,
                    "text": " | ".join(cells),
                    "cells": cells,
                }
            )
            if row_number >= 30000:
                break
    return rows, len(workbook.sheetnames)


def docx_rows(content: bytes) -> tuple[list[dict[str, Any]], int]:
    rows: list[dict[str, Any]] = []
    document = Document(io.BytesIO(content))
    for row_number, paragraph in enumerate(document.paragraphs, 1):
        text = clean(paragraph.text)
        if text:
            rows.append({"location_kind": "docx_paragraph", "row": row_number, "text": text})
    for table_number, table in enumerate(document.tables, 1):
        for row_number, row in enumerate(table.rows, 1):
            cells = [clean(cell.text) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                rows.append(
                    {
                        "location_kind": "docx_table_row",
                        "table": table_number,
                        "row": row_number,
                        "text": " | ".join(cells),
                        "cells": cells,
                    }
                )
    return rows, len(document.tables)


def html_rows(content: bytes, url: str) -> tuple[list[dict[str, Any]], int]:
    soup = BeautifulSoup(content, "html.parser")
    for element in soup(["script", "style", "noscript", "svg"]):
        element.decompose()
    rows: list[dict[str, Any]] = []
    tables = soup.find_all("table")
    for table_number, table in enumerate(tables, 1):
        for row_number, tr in enumerate(table.find_all("tr"), 1):
            cells = [clean(cell.get_text(" ", strip=True)) for cell in tr.find_all(["th", "td"])]
            cells = [cell for cell in cells if cell]
            if cells:
                rows.append(
                    {
                        "location_kind": "html_table_row",
                        "table": table_number,
                        "row": row_number,
                        "text": " | ".join(cells),
                        "cells": cells,
                    }
                )
    main = soup.find("main") or soup.find(id="main") or soup.body or soup
    for line_number, line in enumerate(main.get_text("\n", strip=True).splitlines(), 1):
        text = clean(line)
        if text:
            rows.append(
                {
                    "location_kind": "html_text_line",
                    "row": line_number,
                    "text": text,
                    "url": url,
                }
            )
    return rows, len(tables)


def extract_rows(content: bytes, content_type: str, url: str) -> tuple[list[dict[str, Any]], int]:
    lower_url = url.lower()
    if content_type == "application/pdf" or lower_url.endswith(".pdf"):
        return pdf_rows(content)
    if "spreadsheetml" in content_type or lower_url.endswith((".xlsx", ".xlsm")):
        return workbook_rows(content)
    if "wordprocessingml" in content_type or lower_url.endswith(".docx"):
        return docx_rows(content)
    return html_rows(content, url)


def candidate_rows(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    candidates: list[dict[str, Any]] = []
    seen: set[str] = set()
    for row in rows:
        text = clean(row.get("text"))
        if len(text) < 6 or len(text) > 1200 or not NUMBER_RE.search(text):
            continue
        keywords = [keyword for keyword in TARGET_KEYWORDS if keyword.lower() in text.lower()]
        if not keywords and not row["location_kind"].endswith("table_row"):
            continue
        normalized = re.sub(r"[\s|]+", "", text).lower()
        if normalized in seen:
            continue
        seen.add(normalized)
        candidates.append(
            {
                **{
                    key: value
                    for key, value in row.items()
                    if key in {"location_kind", "page", "sheet", "table", "row", "text"}
                },
                "matched_keywords": keywords,
                "numeric_token_count": len(NUMBER_RE.findall(text)),
            }
        )
    return candidates[:2000]


def audit_prefecture(session: requests.Session, record: dict[str, Any]) -> dict[str, Any]:
    source = record["target_source"]
    result: dict[str, Any] = {
        **{key: value for key, value in record.items() if key != "target_source"},
        "target_source": source,
        "resolution_status": "pending",
        "documents": [],
        "candidate_count": 0,
        "errors": [],
    }
    try:
        documents, landing_audit = resolve_documents(session, source["url"], source["title"])
        result["landing_audit"] = landing_audit
        for document in documents:
            try:
                response = request(session, document.url)
                content_type = normalized_content_type(response) or document.content_type
                rows, structural_count = extract_rows(response.content, content_type, response.url)
                candidates = candidate_rows(rows)
                result["documents"].append(
                    {
                        "url": response.url,
                        "title": document.title,
                        "content_type": content_type,
                        "sha256": sha256_bytes(response.content),
                        "structural_count": structural_count,
                        "row_count": len(rows),
                        "candidate_count": len(candidates),
                        "candidate_sample": candidates[:12],
                    }
                )
                result["candidate_count"] += len(candidates)
            except Exception as exc:  # noqa: BLE001
                result["errors"].append(
                    {"url": document.url, "error": f"{type(exc).__name__}: {exc}"}
                )
        result["resolution_status"] = (
            "resolved_with_candidates"
            if result["candidate_count"]
            else "resolved_without_candidates"
        )
    except Exception as exc:  # noqa: BLE001
        result["resolution_status"] = "failed"
        result["errors"].append({"url": source["url"], "error": f"{type(exc).__name__}: {exc}"})
    return result


def build(root: Path, output: Path) -> None:
    session = requests.Session()
    session.headers.update({"User-Agent": USER_AGENT, "Accept-Language": "ja,en;q=0.6"})
    audited = [audit_prefecture(session, record) for record in official_phase9_records(root)]
    status_counts: dict[str, int] = {}
    for record in audited:
        status_counts[record["resolution_status"]] = status_counts.get(
            record["resolution_status"], 0
        ) + 1
    write_json(
        output,
        {
            "id": "phase9-numeric-target-source-audit",
            "updated_at": UPDATED_AT,
            "prefecture_count": len(audited),
            "document_count": sum(len(record["documents"]) for record in audited),
            "candidate_count": sum(record["candidate_count"] for record in audited),
            "status_counts": status_counts,
            "records": audited,
        },
    )


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--root", type=Path, default=ROOT)
    parser.add_argument(
        "--output",
        type=Path,
        default=ROOT / "data/extracted/phase9_numeric_target_source_audit.json",
    )
    args = parser.parse_args()
    build(args.root, args.output)


if __name__ == "__main__":
    main()
